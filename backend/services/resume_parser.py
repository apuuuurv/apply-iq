"""
Resume file parser service.

This module handles extraction of text from resume files (PDF and DOCX).
It uses:
- pdfplumber: For efficient, accurate PDF text extraction
- python-docx: For reading Microsoft Word documents

The parser handles various edge cases like corrupt files,
unusual formatting, and validates content before returning.
"""

import logging
from typing import Optional
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Handles resume file parsing (PDF and DOCX formats).
    
    This class provides methods to extract text from resume files
    while handling errors gracefully.
    """
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Uses pdfplumber which provides:
        - More accurate text extraction than PyPDF2
        - Better handling of tables and formatting
        - Direct access to page content
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text from the PDF
            
        Raises:
            ValueError: If PDF is corrupted or cannot be read
            IOError: If file doesn't exist
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"Parsing PDF: {file_path} ({len(pdf.pages)} pages)")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        logger.debug(f"Extracted text from page {page_num}")
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {str(e)}")
                        continue
            
            if not text.strip():
                raise ValueError(
                    "No text content found in PDF. "
                    "This may be a scanned/image-based PDF. "
                    "Please use a PDF with selectable text or convert to DOCX format."
                )
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
            
        except FileNotFoundError as e:
            logger.error(f"PDF file not found: {file_path}")
            raise IOError(f"Resume file not found: {file_path}")
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Could not parse PDF file: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from a DOCX (Word) file.
        
        Uses python-docx to:
        - Extract paragraph text
        - Extract table content
        - Handle nested structures
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the DOCX
            
        Raises:
            ValueError: If DOCX is corrupted or cannot be read
            IOError: If file doesn't exist
        """
        try:
            doc = Document(file_path)
            logger.info(f"Parsing DOCX: {file_path}")
            
            text = ""
            
            # Extract from paragraphs using the safe, built-in method
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text += cell.text + " "
                    if row_text.strip():
                        text += row_text + "\n"
            
            if not text.strip():
                raise ValueError("No text content found in DOCX")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
            
        except FileNotFoundError as e:
            logger.error(f"DOCX file not found: {file_path}")
            raise IOError(f"Resume file not found: {file_path}")
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Could not parse DOCX file: {str(e)}")
    
    @staticmethod
    def parse_resume(file_path: str, file_type: str) -> str:
        """
        Parse resume file based on type.
        
        Determines which parser to use based on file extension,
        then returns extracted text.
        
        Args:
            file_path: Path to the resume file
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Extracted resume text
            
        Raises:
            ValueError: If file type is unsupported or file cannot be parsed
        """
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return ResumeParser.parse_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return ResumeParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx")
